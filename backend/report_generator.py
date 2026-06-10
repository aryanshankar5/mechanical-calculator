from docx import Document
from docx2pdf import convert
import os
import subprocess
from datetime import datetime


def replace_placeholders(doc, replacements):

    # Paragraphs
    for paragraph in doc.paragraphs:

        full_text = paragraph.text

        for key, value in replacements.items():
            full_text = full_text.replace(
                "{{" + key + "}}",
                str(value)
            )

        if paragraph.runs:
            paragraph.runs[0].text = full_text

            for run in paragraph.runs[1:]:
                run.text = ""

    # Tables
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    full_text = paragraph.text

                    for key, value in replacements.items():
                        full_text = full_text.replace(
                            "{{" + key + "}}",
                            str(value)
                        )

                    if paragraph.runs:
                        paragraph.runs[0].text = full_text

                        for run in paragraph.runs[1:]:
                            run.text = ""
                                

def generate_bearing_report(data):

    template_path = os.path.join(
        os.getcwd(),
        "templates",
        "Bearing Report.docx"
    )

    output_docx = os.path.join(
        os.getcwd(),
        "reports",
        "bearing_report.docx"
    )

    output_pdf = os.path.join(
        os.getcwd(),
        "reports",
        "bearing_report.pdf"
    )

    doc = Document(template_path)

    print("REPORT DATA")
    for k, v in data.items():
        print(k, "=", v)

    replace_placeholders(
        doc,
        data
    )

    doc.save(output_docx)

    convert(output_docx, output_pdf)

    return output_pdf

def generate_vibrating_screen_report(data):

    template_path = os.path.join(
        os.getcwd(),
        "templates",
        "Vibrating Screen Report.docx"
    )

    output_docx = os.path.join(
        os.getcwd(),
        "reports",
        "vibrating_screen_report.docx"
    )

    output_pdf = os.path.join(
        os.getcwd(),
        "reports",
        "vibrating_screen_report.pdf"
    )

    doc = Document(template_path)

    replace_placeholders(doc, data)

    doc.save(output_docx)

    convert(output_docx, output_pdf)

    return output_pdf

def generate_vibrating_feeder_report(data):

    template_path = os.path.join(
        os.getcwd(),
        "templates",
        "Vibrating Feeder Report.docx"
    )

    output_docx = os.path.join(
        os.getcwd(),
        "reports",
        "vibrating_feeder_report.docx"
    )

    output_pdf = os.path.join(
        os.getcwd(),
        "reports",
        "vibrating_feeder_report.pdf"
    )

    doc = Document(template_path)

    replace_placeholders(doc, data)

    doc.save(output_docx)

    convert(output_docx, output_pdf)

    return output_pdf


def format_value(value):
    if value is None:
        return "-"

    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".")

    return str(value)


def replace_placeholders_in_paragraph(paragraph, data):
    full_text = "".join(run.text for run in paragraph.runs)

    if "{" not in full_text:
        return

    new_text = full_text

    for key, value in data.items():
        placeholder = "{" + key + "}"
        new_text = new_text.replace(placeholder, format_value(value))

    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ""

        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)


def replace_placeholders_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)


def replace_all_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)

    for table in doc.tables:
        replace_placeholders_in_table(table, data)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)

        for table in section.header.tables:
            replace_placeholders_in_table(table, data)

        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)

        for table in section.footer.tables:
            replace_placeholders_in_table(table, data)





def generate_conveyor_report(report_data):
    base_dir = os.path.dirname(os.path.abspath(__file__))

    template_path = os.path.join(
        base_dir,
        "templates",
        "Conveyor final Report.docx"
    )

    output_dir = os.path.join(base_dir, "reports")
    os.makedirs(output_dir, exist_ok=True)

    output_docx = os.path.join(
        output_dir,
        "Conveyor_Report_Filled.docx"
    )

    output_pdf = os.path.join(
        output_dir,
        "Conveyor_Report.pdf"
    )

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    replace_all_placeholders(doc, report_data)

    doc.save(output_docx)

    # Convert DOCX to PDF using docx2pdf
    convert(output_docx, output_pdf)

    if not os.path.exists(output_pdf):
        raise FileNotFoundError(f"PDF was not generated: {output_pdf}")

    return output_pdf